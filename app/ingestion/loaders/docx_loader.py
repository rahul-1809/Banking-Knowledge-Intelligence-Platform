from pathlib import Path

from docx import Document

from app.ingestion.loaders.base import LoadedDocument, build_metadata


def load_docx(file_path: Path, data_root: Path) -> LoadedDocument:
    metadata = build_metadata(file_path, data_root)
    document = Document(str(file_path))
    paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]
    text = "\n\n".join(paragraphs)
    return LoadedDocument(metadata=metadata, text=text)
