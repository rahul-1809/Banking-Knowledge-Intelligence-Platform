"""Generate sample banking documents across 6 categories in multiple file formats (.pdf, .docx, .txt, .md, .json).

Categories:
1. RBI        — Master directions, circulars on KYC, PSL, and capital adequacy.
2. SOP        — Internal SOPs for retail onboarding, EDD, and branch ops.
3. CREDIT     — Personal, SME, and commercial credit underwriting policies.
4. COMPLIANCE — AML/CFT screening, sanctions, and fraud reporting protocols.
5. TREASURY   — Asset Liability Management (ALM), liquidity risk, investment rules.
6. AUDIT      — Internal audit manuals, branch inspection, risk control matrices.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document  # type: ignore[import-untyped]
from reportlab.lib.pagesizes import letter  # type: ignore[import-untyped]
from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import-untyped]
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # type: ignore[import-untyped]


DATA_DIR = Path("DATA")


def create_pdf(filename: Path, title: str, sections: list[tuple[str, str]]) -> None:
    """Generate a clean PDF document using ReportLab."""
    filename.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(str(filename), pagesize=letter)
    styles = getSampleStyleSheet()

    story = []
    story.append(Paragraph(f"<b><font size=16>{title}</font></b>", styles["Title"]))
    story.append(Spacer(1, 12))

    for heading, text in sections:
        story.append(Paragraph(f"<b><font size=12>{heading}</font></b>", styles["Heading2"]))
        story.append(Spacer(1, 4))
        story.append(Paragraph(text, styles["BodyText"]))
        story.append(Spacer(1, 10))

    doc.build(story)


def create_docx(filename: Path, title: str, sections: list[tuple[str, str]]) -> None:
    """Generate a formatted Word document using python-docx."""
    filename.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, level=1)

    for heading, text in sections:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(text)

    doc.save(str(filename))


def create_txt(filename: Path, title: str, sections: list[tuple[str, str]]) -> None:
    """Generate a plain text document."""
    filename.parent.mkdir(parents=True, exist_ok=True)
    content = f"{title}\n\n"
    for heading, text in sections:
        content += f"{heading}\n{text}\n\n"
    filename.write_text(content.strip() + "\n", encoding="utf-8")


def create_json(filename: Path, title: str, metadata: dict, sections: list[tuple[str, str]]) -> None:
    """Generate a structured JSON document."""
    filename.parent.mkdir(parents=True, exist_ok=True)
    data = {
        "title": title,
        "metadata": metadata,
        "sections": [{"heading": h, "body": b} for h, b in sections],
    }
    filename.write_text(json.dumps(data, indent=2), encoding="utf-8")


def generate_all_docs() -> None:
    """Create sample files across all 6 banking document categories."""
    print("Generating multi-format banking documents across 6 categories...")

    # ── Category 1: RBI ──────────────────────────────────────────────────────
    create_pdf(
        DATA_DIR / "RBI" / "rbi_master_direction_kyc.pdf",
        "RBI Master Direction — Know Your Customer (KYC) Norms 2024",
        [
            ("1. Framework & Scope", "Regulated entities (REs) must implement a robust KYC policy covering Customer Acceptance, Risk Categorisation, and Continuous Transaction Monitoring."),
            ("2. Customer Due Diligence (CDD)", "For individual customers, officially valid documents (OVD) accepted include Passport, Driving Licence, Proof of Possession of Aadhaar, Voter Identity Card, and NREGA job card."),
            ("3. Periodic KYC Updation", "KYC records shall be updated at least once every 2 years for high-risk customers, 8 years for medium-risk, and 10 years for low-risk customers."),
        ],
    )

    create_docx(
        DATA_DIR / "RBI" / "rbi_priority_sector_lending.docx",
        "RBI Master Circular — Priority Sector Lending (PSL) Targets",
        [
            ("1. Overall Target", "Domestic commercial banks shall achieve an overall Priority Sector Lending target of 40 percent of Adjusted Net Bank Credit (ANBC) or Credit Equivalent Amount of Off-Balance Sheet Exposure (CEOBE), whichever is higher."),
            ("2. Sub-targets for Agriculture & MSME", "Within the 40 percent target, 18 percent is allocated for Agriculture (with 10 percent for Small & Marginal Farmers) and 7.5 percent for Micro Enterprises."),
            ("3. Non-compliance Penalty", "Shortfalls in achieving PSL targets shall be deposited into the Rural Infrastructure Development Fund (RIDF) established with NABARD or specified funds with SIDBI/NHB."),
        ],
    )

    # ── Category 2: SOP ──────────────────────────────────────────────────────
    create_pdf(
        DATA_DIR / "SOP" / "sop_retail_account_opening.pdf",
        "Internal SOP — Retail Savings & Current Account Opening",
        [
            ("1. Objective & Application", "Standard operating procedure for verifying retail customer applications at branch counters and digital channels."),
            ("2. Verification Steps", "Staff must verify original OVDs against submitted copies, perform In-Person Verification (IPV) or Video KYC (V-CIP), and record customer consent."),
            ("3. Risk Scoring & System Entry", "Assign initial risk rating (Low/Medium/High) in core banking based on customer profile, occupation, and geographic location."),
        ],
    )

    create_docx(
        DATA_DIR / "SOP" / "sop_high_risk_customer_edd.docx",
        "Internal SOP — Enhanced Due Diligence (EDD) Approval Process",
        [
            ("1. Scope", "Mandatory procedure for onboarding Politically Exposed Persons (PEPs), high-net-worth individuals from high-risk jurisdictions, and non-face-to-face accounts."),
            ("2. EDD Requirements", "Staff must obtain verified proof of source of wealth, source of funds, senior management approval, and conduct web screening against sanctions lists."),
            ("3. Escalation & Approval Matrix", "High-risk account activation requires written clearance from the Branch Compliance Officer and Regional AML Head."),
        ],
    )

    # ── Category 3: CREDIT ───────────────────────────────────────────────────
    create_pdf(
        DATA_DIR / "CREDIT" / "credit_policy_sme_loans.pdf",
        "Credit Policy — SME Working Capital & Term Loan Guidelines",
        [
            ("1. Borrower Eligibility", "Micro, Small, and Medium Enterprises (MSMEs) with a minimum operating track record of 3 years, audited balance sheets, and a Minimum Debt Service Coverage Ratio (DSCR) of 1.35x."),
            ("2. Collateral & Security", "Primary security: Hypothecation of stocks, receivables, and plant/machinery. Collateral coverage ratio must be at least 125% of loan exposure for unassisted loans."),
            ("3. Financial Covenants", "Minimum Current Ratio of 1.25x and Maximum Total Outside Liabilities to Tangible Net Worth (TOL/TNW) of 3.0x must be maintained throughout the tenure."),
        ],
    )

    create_docx(
        DATA_DIR / "CREDIT" / "credit_policy_retail_mortgage.docx",
        "Credit Policy — Home Loan & LAP Underwriting Manual",
        [
            ("1. Loan-to-Value (LTV) Limits", "For home loans up to INR 30 Lakhs, maximum LTV ratio is 90%. For loans between INR 30 Lakhs and 75 Lakhs, maximum LTV is 80%. For loans above 75 Lakhs, maximum LTV is 75%."),
            ("2. Fixed Obligation to Income Ratio (FOIR)", "Maximum allowable FOIR including proposed housing EMI is capped at 55% for salaried applicants and 60% for self-employed professionals."),
            ("3. Title Legal Clearance", "Mandatory 30-year search report from bank-empanelled advocate and technical valuation report by two independent valuers for property value exceeding INR 1 Crore."),
        ],
    )

    # ── Category 4: COMPLIANCE ───────────────────────────────────────────────
    create_pdf(
        DATA_DIR / "COMPLIANCE" / "aml_sanctions_screening_policy.pdf",
        "Compliance Policy — Anti-Money Laundering & Sanctions Screening",
        [
            ("1. Policy Statement", "The Bank is committed to zero tolerance for money laundering, terrorist financing, and proliferation financing across all operating jurisdictions."),
            ("2. Sanctions Screening", "All outward and inward cross-border wire transfers, trade finance transactions, and new customer names must be screened real-time against UN Security Council, OFAC, EU, and domestic sanctions lists."),
            ("3. Suspicious Transaction Reporting (STR)", "Transactions flagged for unusual patterns, structurings, or red flags must be analyzed by the AML Monitoring Cell and reported to FIU-IND within 7 working days of confirmation."),
        ],
    )

    create_docx(
        DATA_DIR / "COMPLIANCE" / "fraud_prevention_escalation.docx",
        "Compliance Framework — Internal & External Fraud Prevention",
        [
            ("1. Early Warning Signals (EWS)", "Branches must monitor red flag indicators such as frequent cheque bounces, unexplained cash deposits, unauthorized system logons, and unexpected collateral substitutions."),
            ("2. Incident Escalation Timeline", "Any confirmed or suspected fraud exceeding INR 10 Lakhs must be reported to the Chief Vigilance Officer (CVO) within 24 hours and to the RBI Fraud Monitoring Cell within 3 weeks."),
            ("3. Whistleblower Protection", "Employees reporting suspicious activities in good faith are protected from retaliation under the Bank's Whistleblower Policy."),
        ],
    )

    # ── Category 5: TREASURY ─────────────────────────────────────────────────
    create_pdf(
        DATA_DIR / "TREASURY" / "treasury_alm_liquidity_risk.pdf",
        "Treasury Policy — Asset Liability Management & Liquidity Risk",
        [
            ("1. Liquidity Coverage Ratio (LCR)", "The Bank shall maintain a minimum Liquidity Coverage Ratio (LCR) of 100 percent on an ongoing basis, consisting of unencumbered High Quality Liquid Assets (HQLA)."),
            ("2. Structural Liquidity Gaps", "Cumulative negative mismatches in time buckets up to 14 days shall not exceed 10 percent of cumulative cash outflows in those buckets."),
            ("3. Interest Rate Risk in Banking Book (IRRBB)", "Impact of a +/- 200 basis point parallel interest rate shock on Net Interest Income (NII) shall not exceed 10 percent of annual NII."),
        ],
    )

    create_docx(
        DATA_DIR / "TREASURY" / "treasury_investment_guidelines.docx",
        "Treasury Policy — SLR & Non-SLR Investment Operations",
        [
            ("1. Statutory Liquidity Ratio (SLR)", "Maintain prescribed SLR investments in approved Central and State Government securities as mandated by RBI under Section 24 of the Banking Regulation Act."),
            ("2. Non-SLR Investments", "Investments in corporate bonds and commercial papers are restricted to AAA or AA+ rated instruments. Exposure to any single corporate group is capped at 15 percent of capital funds."),
            ("3. Mark-to-Market (MTM) Valuation", "Held for Trading (HFT) and Available for Sale (AFS) securities must be revalued on a weekly and monthly basis respectively per RBI valuation norms."),
        ],
    )

    create_json(
        DATA_DIR / "TREASURY" / "treasury_investment_limits.json",
        "Treasury Operations — Counterparty & Dealer Exposure Limits",
        {"category": "TREASURY", "effective_date": "2024-04-01"},
        [
            ("Interbank Money Market Limit", "Maximum overnight call/notice money borrowing limit is capped at 100% of capital funds on a fortnight average basis."),
            ("Derivative Counterparty Exposure", "Over-the-Counter (OTC) interest rate swap exposure to non-bank counterparties requires Credit Support Annex (CSA) collateral agreements."),
        ],
    )

    # ── Category 6: AUDIT ────────────────────────────────────────────────────
    create_pdf(
        DATA_DIR / "AUDIT" / "internal_audit_manual.pdf",
        "Internal Audit Policy — Risk-Based Branch Inspection Manual",
        [
            ("1. Audit Frequency & Rating", "Branches are audited under Risk-Based Internal Audit (RBIA) methodology every 12 to 18 months based on operational risk categorization (High, Medium, Low)."),
            ("2. Key Checkpoints", "Inspect cash vault double-custody records, loan file security documentation, dormant account reactivation approvals, and KYC compliance samples."),
            ("3. Compliance & Rectification", "Branch managers must rectify high-risk audit observations within 30 days of report submission. Unresolved items are escalated to the Audit Committee of the Board (ACB)."),
        ],
    )

    create_docx(
        DATA_DIR / "AUDIT" / "operational_risk_matrix.docx",
        "Operational Risk & Internal Control Assessment Matrix",
        [
            ("1. Risk Assessment Framework", "Identifies Key Risk Indicators (KRIs) across branch operations, IT systems, clearing operations, and credit administration."),
            ("2. Loss Event Data Collection", "All operational risk losses exceeding INR 50,000 must be recorded in the Bank's Loss Event Database within 5 working days."),
            ("3. Control Effectiveness Testing", "Quarterly testing of key operational controls by internal risk managers to ensure compliance with Basel III Operational Risk guidelines."),
        ],
    )

    create_txt(
        DATA_DIR / "AUDIT" / "branch_inspection_checklist.txt",
        "Internal Audit — Branch Operations Quick Inspection Checklist",
        [
            ("Vault & Cash Management", "Verify dual key ownership, daily cash register reconciliation, and alarm system testing logs."),
            ("Dormant Accounts", "Ensure customer presence or written request along with fresh KYC documentation before reactivating dormant accounts."),
        ],
    )

    print("Successfully generated files for all 6 categories!")


if __name__ == "__main__":
    generate_all_docs()
